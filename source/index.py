import time
from os.path import split

from PyQt5 import QtWidgets, QtGui, QtCore, uic
from PyQt5.QtCore import QFileInfo
from PyQt5.QtWidgets import QLabel
from PyQt5.QtChart import *
from PyQt5.Qt import Qt
from PyQt5 import QtChart
# from PyQt5.QtChart import *
from PyQt5.QtGui import QPainter, QIntValidator, QPixmap, QImage

import sys, os, random, threading, platform
import pandas as pd
import logging
from plyer import notification
from PyQt5.QtWidgets import QHeaderView, QTableWidgetItem
from pandas.tests.io.excel.test_xlrd import xlwt
import matplotlib
import matplotlib.pyplot as plt
import numpy as np

from docx import Document  # python-docx
from docx.shared import Inches

DESKTOP = os.path.join(os.path.join(os.environ['USERPROFILE']),
                       'Desktop') if platform.system() == 'Windows' else os.path.join(
    os.path.join(os.path.expanduser('~')), 'Desktop')


m_e = ["Electrical Machines and Equipment's", "Electrical Power", "Electric ground lines", "Electrician distribute", "Electric overhead lines", "Mechanical Production", "Refrigeration and Air Conditioning", "Automotive Mechanics", "Computer Technical Support", "Office Management", "Accounting", "Marketing", "Industrial Electronics and Control", "Biomedical Technology", "Biomedical Technology", "Architectural construction", "Civilization Construction", "Surveying", "Chemical Production", "Chemical laboratories", "Food Safety", "Environmental Protection"]


class Splash(QtWidgets.QWidget):
    def __init__(self):
        super().__init__()
        uic.loadUi(os.path.join(os.path.dirname(__file__), 'ui/splash.ui'), self)
        print(os.path.join(os.path.dirname(__file__), 'ui/splash.ui'))
        self.en.installEventFilter(self)
        self.ar.installEventFilter(self)

    def eventFilter(self, s, e):
        if e.type() == QtCore.QEvent.MouseButtonPress:
            if s is self.en:
                self.main = Main()
                self.main.show()
                self.close()
            if s is self.ar:
                self.main = Main(ar=True)
                self.main.show()
                self.close()
        return super(Splash, self).eventFilter(s, e)


# todo############################################### Main window #######################################################
class Main(QtWidgets.QWidget):
    def __init__(self, ar=False):
        super().__init__()
        logging.info('app started')
        uic.loadUi(os.path.join(os.path.dirname(__file__), 'ui/home.ui'), self)
        self.ar = ar
        ttl = 'الصفحة الرئسية' if self.ar else 'Main'

        self.setWindowTitle(ttl)
        self.browse.clicked.connect(self.get_path)
        self.paths = []
        self.current_R = 0
        self.current_F = ''
        self.proc.setEnabled(False)
        self.path_txt.currentTextChanged.connect(self.path_changed)
        self.loading = Loading('proc.gif')
        self.contents.addWidget(self.loading)
        self.proc.clicked.connect(self.start_proc)
        self.frame_2.setEnabled(False)
        self.r1_btn.installEventFilter(self)
        self.r2_btn.installEventFilter(self)
        self.r3_btn.installEventFilter(self)
        self.r4_btn.installEventFilter(self)
        self.r5_btn.installEventFilter(self)
        self.r6_btn.installEventFilter(self)
        if self.ar:
            self.frame_2.setLayoutDirection(QtCore.Qt.RightToLeft)
            self.path_txt.clear()
            self.path_txt.addItems(['إختر ملف ...'])
            self.browse.setText('تصفح')
            self.proc.setText('إبدأ المعالجة')
            self.r1_btn.setText('التقرير الأول')
            self.r2_btn.setText('التقرير الثاني')
            self.r3_btn.setText('التقرير الثالت')
            self.r4_btn.setText('التقرير الرابع')
            self.r5_btn.setText('التقرير الخامس')
            self.r6_btn.setText('التقرير السادس')

    def start_proc(self):
        if self.path_txt.currentText() != self.current_F:
            self.proc.setEnabled(False)

            self.df = pd.read_excel(self.path_txt.currentText())
            cols = ['Student ID', 'Date of birth', 'Place of birth', 'Type of ID',
                    'Place of issue', 'Department', 'Major', 'Graduation Year', 'Year',
                    'Semester for graduation', 'GPA', 'Grade', 'Type of certificate']
            if len(self.df.columns) != 13:
                notification.notify(title='حدث خطأ مع الملف' if self.ar else 'Error found while looading',
                                    message='المرجو التأكد من الملف أولا' if self.ar else 'make sure you are using the right file',
                                    timeout=7)
                self.clear_content()
                # self.contents.addWidget(Loading('err.gif'))
                self.contents.addWidget(Err())
                return

            # print(self.df)

            start_year = []
            for i in self.df['Student ID']:
                if str(i)[0] == '1':
                    start_year.append(int(f'14{str(i)[3:5]}'))
                elif str(i)[0] == '4':
                    start_year.append(int(f'14{str(i)[1:3]}'))

            self.df['start_year'] = start_year

            self.df['year_in_college'] = self.df['Graduation Year'] - self.df['start_year']
            self.clear_content()
            self.r1 = R1(self.df, self.ar)
            self.contents.addWidget(self.r1)
            self.frame_2.setEnabled(True)
            self.clicks_btns(self.r1_btn)
            self.current_R = 1
            self.current_F = self.path_txt.currentText()
            self.r2 = R2(self.df, self.ar)
            self.r3 = R3(self.df, self.ar)
            self.r4 = R4(self.df, self.ar)
            self.r5 = R5(self.df, self.ar)
            self.r6 = R6(self.df, self.ar)

    def eventFilter(self, o, e):
        if e.type() == QtCore.QEvent.MouseButtonPress or e.type() == QtCore.QEvent.MouseButtonDblClick:
            if o is self.r1_btn and self.current_R != 1:
                self.clicks_btns(self.r1_btn)
                self.current_R = 1
                self.clear_content()
                self.contents.addWidget(self.r1)
                self.setWindowTitle('التقرير الأول' if self.ar else 'Rapport 1')

            elif o is self.r2_btn and self.current_R != 2:
                self.clicks_btns(self.r2_btn)
                self.current_R = 2
                self.clear_content()
                self.contents.addWidget(self.r2)
                self.setWindowTitle('التقرير الثاني' if self.ar else 'Rapport 2')

            elif o is self.r3_btn and self.current_R != 3:
                self.clicks_btns(self.r3_btn)
                self.current_R = 3
                self.clear_content()
                self.contents.addWidget(self.r3)
                self.setWindowTitle('التقرير الثالت' if self.ar else 'Rapport 3')

            elif o is self.r4_btn and self.current_R != 4:
                self.clicks_btns(self.r4_btn)
                self.current_R = 4
                self.clear_content()
                self.contents.addWidget(self.r4)
                self.setWindowTitle('التقرير الرابع' if self.ar else 'Rapport 4')


            elif o is self.r5_btn and self.current_R != 5:
                self.clicks_btns(self.r5_btn)
                self.current_R = 5
                self.clear_content()
                self.contents.addWidget(self.r5)
                self.setWindowTitle('التقرير الخامس' if self.ar else 'Rapport 5')


            elif o is self.r6_btn and self.current_R != 6:
                self.clicks_btns(self.r6_btn)
                self.current_R = 6
                self.clear_content()
                self.contents.addWidget(self.r6)
                self.setWindowTitle('التقرير السادس' if self.ar else 'Rapport 6')

        return super(Main, self).eventFilter(o, e)

    def clicks_btns(self, btn):
        self.r1_btn.setStyleSheet('QPushButton{background-color:white;}')
        self.r2_btn.setStyleSheet('QPushButton{background-color:white;}')
        self.r3_btn.setStyleSheet('QPushButton{background-color:white;}')
        self.r4_btn.setStyleSheet('QPushButton{background-color:white;}')
        self.r5_btn.setStyleSheet('QPushButton{background-color:white;}')
        self.r6_btn.setStyleSheet('QPushButton{background-color:white;}')
        btn.setStyleSheet('QPushButton{background-color:#e6ffe6;}')

    def path_changed(self):
        if self.path_txt.currentText() == '':
            self.path_label.setFixedWidth(0)
            self.proc.setEnabled(False)

        elif os.path.isfile(self.path_txt.currentText()):
            self.path_label.setPixmap(QtGui.QPixmap('src/approved.png'))
            self.path_label.setScaledContents(True)
            self.proc.setEnabled(True)
            self.path_label.setFixedWidth(38)
        else:
            self.path_label.setPixmap(QtGui.QPixmap('src/err.png'))
            self.path_label.setScaledContents(True)
            self.proc.setEnabled(False)
            self.path_label.setFixedWidth(38)

    def get_path(self):
        self.file = QtWidgets.QFileDialog.getOpenFileName(caption='أختيار الملف' if self.ar else 'Load File',
                                                          filter="Excel (*.xlsx *.xls)", directory=DESKTOP)[0]
        if self.file:
            self.proc.setEnabled(True)
            self.paths.insert(0, self.file)
            self.path_txt.clear()
            self.path_txt.addItems(set(self.paths))
            self.path_txt.setCurrentText(self.paths[0])

    def clear_content(self):
        for i in reversed(range(self.contents.count())):
            self.contents.itemAt(i).widget().setParent(None)
        # self.content.addWidget(widget)


# todo############################################### Rapport 1 #########################################################
class R1(QtWidgets.QWidget):
    def __init__(self, df=None, ar=False):
        super().__init__()
        uic.loadUi(os.path.join(os.path.dirname(__file__), 'ui/r1.ui'), self)
        self.df = df
        self.ar = ar
        # self.table.itemSelectionChanged.connect(self.table_select_event)
        self.table.clear()
        self.table_header = ['التخصصات' if self.ar else 'Majors', 'أقل من سنتان' if self.ar else 'less then 2 Years',
                             'سنتان' if self.ar else '2 Years', '3 سنوات' if self.ar else '3 Years',
                             '4 سنوات' if self.ar else '4 Years', '5 سموات و أكثر' if self.ar else '5 Yrs or more']
        self.table.setColumnCount(len(self.table_header))
        self.table.setHorizontalHeaderLabels(self.table_header)
        self.table.resizeColumnsToContents()
        self.from_txt.setValidator(QIntValidator())
        self.to_txt.setValidator(QIntValidator())
        for i in range(len(self.table_header)):
            self.table.horizontalHeader().setSectionResizeMode(i, QHeaderView.Stretch)

        self.set_dt()
        self.filter.clicked.connect(self.filtering)
        self.produce.clicked.connect(self.export_to_exel)
        if self.ar:
            self.frame.setLayoutDirection(QtCore.Qt.RightToLeft)
            self.label_2.setText('من')
            self.label_3.setText('إلى')
            self.from_txt.setPlaceholderText('سنة')
            self.to_txt.setPlaceholderText('سنة')
            self.filter.setText('فلتر')
            self.produce.setText('إستخراج')

    def export_to_exel(self):
        # if self.table.rowCount():
        #     filename = QtWidgets.QFileDialog.getSaveFileName(caption='إستخراج' if self.ar else 'Export', filter="Excel (*.xlsx *.xls)", directory= DESKTOP)[0]
        #     if not QFileInfo(filename).suffix():
        #         filename += '.xlsx'
        #
        #     wbk = xlwt.Workbook()
        #     sheet = wbk.add_sheet("sheet", cell_overwrite_ok=True)
        #     for i, v in enumerate(self.table_header):
        #         sheet.write(0, i, v)
        #     for currentColumn in range(self.table.columnCount()):
        #         for currentRow in range(self.table.rowCount()):
        #             try:
        #                 teext = str(self.table.item(currentRow, currentColumn).text())
        #                 sheet.write(currentRow + 1, currentColumn, teext)
        #             except AttributeError:
        #                 pass
        #
        #     wbk.save(filename)
        #     if filename.split('.')[0]:notification.notify(title= 'تم حفظ الملف بنجاح' if self.ar else 'File Saved Successfully', message=f'Saved at : {filename}',timeout=5)
        # else:
        #     msg = '<font color="red">خطأ : </font>لا توجد تيانات للإستخراج' if self.ar else '<font color="red">ERROR : </font>No Data In The Table To Export'
        #     self.err.setText()

        # print(self.chartView.winId())
        # # pos = self.graph_layout.pos()
        # print(self.chartView.width())
        # scr = self.grabWindow(self.chartView.winId(), 0, 0, self.chartView.width(), self.chartView.height())
        # scr.save('tttt.png')
        # print('done')
        # d = Drawing(280, 250)
        # bar = VerticalBarChart()
        # bar.x = 50
        # bar.y = 85
        #
        # print(data)
        # bar.data = data
        # bar.categoryAxis.categoryNames = lbls
        # # bar.bars[0].fillColor = PCMYKColor(0, 100, 100, 40, alpha=85)
        # # bar.bars[1].fillColor = PCMYKColor(23, 51, 0, 4, alpha=85)
        # # bar.bars.fillColor = PCMYKColor(100, 0, 90, 50, alpha=85)
        # d.add(bar, '')
        # # d.save(formats=['pdf'], outDir='.', fnRoot='test')
        # d.save('test.pdf')

        #
        # x = np.arange(len([i[0] for i in data]))  # the label locations
        # width = 0.2  # the width of the bars
        #
        # fig, ax = plt.subplots()
        # rects1 = ax.bar(x - width / 2, [i[1] for i in data], width, label='2')
        # rects2 = ax.bar(x + width / 2, [i[2] for i in data], width, label='3')
        # rects3 = ax.bar(x + width / 2, [i[3] for i in data], width, label='4')
        # rects4 = ax.bar(x + width / 2, [i[4] for i in data], width, label='5')
        #
        # # Add some text for labels, title and custom x-axis tick labels, etc.
        # ax.set_ylabel('Scores')
        # ax.set_title('Scores by group and gender')
        # ax.set_xticks(x)
        # ax.set_xticklabels(lbls)
        # ax.legend()
        #
        # def autolabel(rects):
        #     """Attach a text label above each bar in *rects*, displaying its height."""
        #     for rect in rects:
        #         height = rect.get_height()
        #         ax.annotate('{}'.format(height),
        #                     xy=(rect.get_x() + rect.get_width() / 2, height),
        #                     xytext=(0, 3),  # 3 points vertical offset
        #                     textcoords="offset points",
        #                     ha='center', va='bottom')
        #
        # autolabel(rects1)
        # autolabel(rects2)
        # autolabel(rects3)
        # autolabel(rects4)
        #
        # fig.tight_layout()

        data = []
        for r in range(self.table.rowCount()):
            rr = []
            for c in range(self.table.columnCount()):
                rr.append(self.table.item(r, c).text() if self.table.item(r, c).text() else 0)
            data.append(rr)

        import arabic_reshaper
        from bidi.algorithm import get_display  # python-bidi
        plotdata = pd.DataFrame({
            get_display(arabic_reshaper.reshape('أقل من سنتان')) if self.ar else 'less then 2 Years': [int(i[1]) for i in data],
            get_display(arabic_reshaper.reshape('سنتان' )) if self.ar else '2 Years': [int(i[2]) for i in data],
            get_display(arabic_reshaper.reshape('3 سنوات')) if self.ar else '3 Years': [int(i[3]) for i in data],
            get_display(arabic_reshaper.reshape('4 سنوات')) if self.ar else '4 Years': [int(i[4]) for i in data],
            get_display(arabic_reshaper.reshape('5 سنوات و أكثر')) if self.ar else '5 Yrs or more': [int(i[5]) for i in data]
        },
            index=[get_display(arabic_reshaper.reshape(i[0])) for i in data]
        )

        plotdata.plot(kind="bar", figsize=(12, 7))
        plt.title(get_display(arabic_reshaper.reshape(self.title_.text())) if self.ar else self.title_.text())
        # plt.xlabel("Family Member")
        # plt.ylabel("Pies Consumed")
        plt.xticks(rotation=45, horizontalalignment='right')

        grapg_path = os.path.join(DESKTOP, 'graph.png')
        plt.savefig(grapg_path)

        # plt.show()
        document = Document()

        document.add_paragraph(self.title_.text(),
            style='Intense Quote')

        document.add_picture(grapg_path, width=Inches(7))
        table = document.add_table(rows=1, cols=6)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "  "
        hdr_cells[1].text = 'أقل من سنتان' if self.ar else 'less then 2 Years'
        hdr_cells[2].text = 'سنتان' if self.ar else '2 Years'
        hdr_cells[3].text = '3 سنوات' if self.ar else '3 Years'
        hdr_cells[4].text = '4 سنوات' if self.ar else '4 Years'
        hdr_cells[5].text = '5 سنوات و أكثر' if self.ar else '5 Yrs or more'

        for i in data:
            row_cells = table.add_row().cells
            for c in range(6):
                row_cells[c].text = str(i[c])

        if self.table.rowCount():
            self.filename = QtWidgets.QFileDialog.getSaveFileName(caption='إستخراج' if self.ar else 'Export',
                                                                  filter="Word (*.doc *.docx)", directory=DESKTOP)[0]
            if not QFileInfo(self.filename).suffix():
                self.filename += '.docx'

        if self.filename:
            document.save(self.filename)
        print('done')
        os.remove(grapg_path)

    def filtering(self):
        if self.from_txt.text() and self.to_txt.text():
            self.set_dt(int(self.from_txt.text()), int(self.to_txt.text()))
        else:
            # self.err.setText('<font color="red">ERROR : </font>you have to fill from -> to Graduation Year for filtering ')
            self.set_dt()

    def set_dt(self, from_=None, to=None):
        G_years = list([i for i in self.df['Graduation Year']])
        if from_ is None or from_ < min(G_years) or from_ > max(G_years):
            from_ = min(G_years)
        if to is None or to > max(G_years) or to < min(G_years):
            to = max(G_years)
        msg_ar = 'عدد السنوات في المدرسة/الكلية حسب التخصصات من سنة التخرج '
        self.title_.setText(
            msg_ar if self.ar else f'Number of Year in College Grouping by Major from Graduation Year "{from_}" to "{to}"')
        self.new_df = self.df[(self.df['Graduation Year'] >= from_) & (self.df['Graduation Year'] <= to)]
        rows = len(self.new_df)
        if rows > 0:
            majors = [i for i in set(list(self.new_df['Major']))]
            gk = self.new_df.groupby('Major')
            self.data = []
            for major in majors:
                l2 = 0
                _2y = 0
                _3y = 0
                _4y = 0
                _5ym = 0
                for yoc in gk.get_group(major)['year_in_college']:
                    if yoc < 2:
                        l2 += 1
                    elif yoc == 2:
                        _2y += 1
                    elif yoc == 3:
                        _3y += 1
                    elif yoc == 4:
                        _4y += 1
                    elif yoc >= 5:
                        _5ym += 1

                # data.append({major : [_2y, _3y, _4y, _5ym]})
                self.data.append([major, l2, _2y, _3y, _4y, _5ym])
            # [print(i) for i in data]
            [self.table.removeRow(0) for _ in range(self.table.rowCount())]
            for r_n, r_d in enumerate(self.data):
                self.table.insertRow(r_n)
                for c_n, d in enumerate(r_d):
                    self.table.setItem(r_n, c_n, QtWidgets.QTableWidgetItem(str(d)))

            sett = QtChart.QBarSet('أقل من سنتان' if self.ar else "less than 2 year")
            set0 = QtChart.QBarSet('سنتان' if self.ar else "2 year")
            set1 = QtChart.QBarSet('3 سنوات' if self.ar else "3 year")
            set2 = QtChart.QBarSet('4 سنوات' if self.ar else "4 year")
            set3 = QtChart.QBarSet('5 سموات و أكثر' if self.ar else "5 year")

            sett.append([i[1] for i in self.data])
            set0.append([i[2] for i in self.data])
            set1.append([i[3] for i in self.data])
            set2.append([i[4] for i in self.data])
            set3.append([i[5] for i in self.data])

            series = QtChart.QBarSeries()
            series.append(sett)
            series.append(set0)
            series.append(set1)
            series.append(set2)
            series.append(set3)

            chart = QtChart.QChart()
            chart.addSeries(series)
            chart.setTitle(' ')
            chart.setAnimationOptions(QtChart.QChart.SeriesAnimations)

            axisX = QtChart.QBarCategoryAxis()
            # axisX.append(str(i) for i in range(1, len(data) + 1))
            axisX.append(str(i[0]) for i in self.data)
            axisX.setLabelsAngle(90)
            axisX.setTitleText('التخصصات' if self.ar else "Majors")
            font = QtGui.QFont()
            font.setPixelSize(5)
            axisX.tickFont = font

            all_v = []
            for i in self.data:
                for x in i[1:]:
                    all_v.append(x)
            axisY = QtChart.QValueAxis()
            axisY.setRange(0, max(all_v) if all_v else 0)
            axisY.setTitleText('عدد السنوات في الكلية' if self.ar else "Years in College")

            chart.addAxis(axisX, Qt.AlignBottom)
            chart.addAxis(axisY, Qt.AlignLeft)

            chart.legend().setVisible(True)
            chart.legend().setAlignment(Qt.AlignBottom)
            self.chartView = QtChart.QChartView(chart)

            for i in reversed(range(self.verticalLayout_3.count())):
                self.verticalLayout_3.itemAt(i).widget().setParent(None)

            self.verticalLayout_3.addWidget(self.frame_2)
            for i in reversed(range(self.graph_layout.count())):
                self.graph_layout.itemAt(i).widget().setParent(None)

            self.graph_layout.addWidget(self.chartView)
        else:
            lbl = Err()
            for i in reversed(range(self.verticalLayout_3.count())):
                self.verticalLayout_3.itemAt(i).widget().setParent(None)
            self.verticalLayout_3.addWidget(lbl)


# todo############################################### Rapport 2 #########################################################
class R2(QtWidgets.QWidget):
    def __init__(self, df=None, ar=False):
        super().__init__()
        uic.loadUi(os.path.join(os.path.dirname(__file__), 'ui/r2.ui'), self)
        self.df = df
        self.ar = ar
        # self.table.itemSelectionChanged.connect(self.table_select_event)
        self.table.clear()

        self.table.resizeColumnsToContents()
        self.from_txt.setValidator(QIntValidator())
        self.to_txt.setValidator(QIntValidator())
        for i in range(4):
            self.table.horizontalHeader().setSectionResizeMode(i, QHeaderView.Stretch)

        self.filter.clicked.connect(self.filtering)
        # self.filter.clicked.connect(self.filtering)
        self.produce.clicked.connect(self.export_to_exel)
        if self.ar:
            self.frame.setLayoutDirection(QtCore.Qt.RightToLeft)
            self.label_2.setText('من')
            self.label_3.setText('إلى')
            self.from_txt.setPlaceholderText('سنة')
            self.to_txt.setPlaceholderText('سنة')
            self.filter.setText('فلتر')
            self.produce.setText('إستخراج')
            self.label.setText("تصنيف حسب")
            self.comboBox.clear()
            self.comboBox.addItems(["الكليات", "التخصصات", "الأقسام"])

        self.set_dt()

    def export_to_exel(self):
        if self.table.rowCount():

            data = []
            for r in range(self.table.rowCount()):
                rr = []
                for c in range(self.table.columnCount()):
                    rr.append(self.table.item(r, c).text() if self.table.item(r, c).text() else 0)
                data.append(rr)

            import arabic_reshaper
            from bidi.algorithm import get_display  # python-bidi

            self.plotdata = None

            self.plotdata = pd.DataFrame({
                get_display(arabic_reshaper.reshape("على الأقل")) if self.ar else "min": [int(i[1]) for i in data],
                get_display(arabic_reshaper.reshape("المتوسط")) if self.ar else "mean": [int(i[2]) for i in data],
                get_display(arabic_reshaper.reshape("على الأكثر")) if self.ar else "max": [int(i[3]) for i in data]},
                index=[get_display(arabic_reshaper.reshape(i[0])) for i in data]
            )

            self.plotdata.plot(kind="bar", figsize=(12, 7))
            plt.title(get_display(arabic_reshaper.reshape(self.title.text())) if self.ar else self.title.text())
            # plt.xlabel("Family Member")
            # plt.ylabel("Pies Consumed")
            plt.xticks(rotation=45, horizontalalignment='right')

            grapg_path = os.path.join(DESKTOP, 'graph.png')
            plt.savefig(grapg_path)

            # plt.show()
            document = Document()

            document.add_paragraph(self.title.text(),
                style='Intense Quote')

            document.add_picture(grapg_path, width=Inches(7))
            table = document.add_table(rows=1, cols=4)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = " - "
            hdr_cells[1].text = "على الأقل" if self.ar else "min"
            hdr_cells[2].text = "المتوسط" if self.ar else "mean"
            hdr_cells[3].text = "على الأكثر" if self.ar else "max"

            for i in data:
                row_cells = table.add_row().cells
                for c in range(4):
                    row_cells[c].text = str(i[c])

            if self.table.rowCount():
                self.filename = QtWidgets.QFileDialog.getSaveFileName(caption='إستخراج' if self.ar else 'Export',
                                                                      filter="Word (*.doc *.docx)", directory=DESKTOP)[
                    0]
                if not QFileInfo(self.filename).suffix():
                    self.filename += '.docx'

            if self.filename:
                document.save(self.filename)
            print('done')
            os.remove(grapg_path)

    def set_dt(self, from_=None, to=None):
        G_years = list([i for i in self.df['Graduation Year']])
        if from_ is None or from_ < min(G_years) or from_ > max(G_years):
            from_ = min(G_years)
        if to is None or to > max(G_years) or to < min(G_years):
            to = max(G_years)
        self.headers = []
        self.rows = []
        self.dff = self.df[(self.df['Graduation Year'] >= from_) & (self.df['Graduation Year'] <= to)]

        if self.comboBox.currentIndex() == 0:
            ll = list([i for i in self.dff['Graduation Year'] if i])
            self.rows = [['الكلية' if self.ar else 'College', min(ll), int(sum(ll) / len(ll)), max(ll)]]
            self.headers = ['الكليات' if self.ar else 'Colleges', 'على الأقل' if self.ar else 'Min',
                            'المتوسط' if self.ar else 'Mean', 'على الأكثر' if self.ar else 'Max']

        else:
            key = self.comboBox.currentText()
            if self.ar:
                if self.comboBox.currentIndex() == 1:
                    key = 'Major'
                elif self.comboBox.currentIndex() == 2:
                    key = 'Department'

            elements = [i for i in set(list(self.dff[key]))]
            self.data = self.dff.groupby(key)
            self.headers = [self.comboBox.currentText(),'على الأقل' if self.ar else 'Min',
                            'المتوسط' if self.ar else 'Mean', 'على الأكثر' if self.ar else 'Max']
            for elm in elements:
                ll = [i for i in self.data.get_group(elm)['Graduation Year']]
                self.rows.append([elm, min(ll), int(sum(ll) / len(ll)), max(ll)])
        msg = 'تصنيف عدد السنوات في الكلية' if self.ar else f'min/mean/max of Graduation Year Grouping by {self.comboBox.currentText()}s from Graduation Year "{from_}" to "{to}"'
        self.title.setText(f'{msg}')
        self.table.setHorizontalHeaderLabels(self.headers)
        self.table.setColumnCount(len(self.headers))
        [self.table.removeRow(0) for _ in range(self.table.rowCount())]
        for r_n, r_d in enumerate(self.rows):
            self.table.insertRow(r_n)
            for c_n, d in enumerate(r_d):
                self.table.setItem(r_n, c_n, QtWidgets.QTableWidgetItem(str(d)))

        set0 = QtChart.QBarSet('على الأقل')
        set1 = QtChart.QBarSet('المتوصط')
        set2 = QtChart.QBarSet('على الأكثر')

        set0.append([i[1] for i in self.rows])
        set1.append([i[2] for i in self.rows])
        set2.append([i[3] for i in self.rows])

        series = QtChart.QBarSeries()
        series.append(set0)
        series.append(set1)
        series.append(set2)

        chart = QtChart.QChart()
        chart.addSeries(series)
        chart.setTitle(' ')
        chart.setAnimationOptions(QtChart.QChart.SeriesAnimations)

        axisX = QtChart.QBarCategoryAxis()
        # axisX.append(str(i) for i in range(1, len(data) + 1))
        axisX.append(str(i[0]) for i in self.rows)
        axisX.setLabelsAngle(90)
        axisX.setTitleText(f'{self.comboBox.currentText()}')
        font = QtGui.QFont()
        font.setPixelSize(5)
        axisX.tickFont = font
        all_v = []
        for i in self.rows:
            for x in i[1:]:
                all_v.append(x)
        axisY = QtChart.QValueAxis()
        axisY.setRange(0, max(all_v) if all_v else 0)
        axisY.setTitleText("سنوات التخرج" if self.ar else "Graduation Year")

        chart.addAxis(axisX, Qt.AlignBottom)
        chart.addAxis(axisY, Qt.AlignLeft)

        chart.legend().setVisible(True)
        chart.legend().setAlignment(Qt.AlignBottom)
        chartView = QtChart.QChartView(chart)

        for i in reversed(range(self.graph_layout.count())):
            self.graph_layout.itemAt(i).widget().setParent(None)

        self.graph_layout.addWidget(chartView)

    def filtering(self):
        if self.from_txt.text() and self.to_txt.text():
            self.set_dt(int(self.from_txt.text()), int(self.to_txt.text()))
        else:
            # self.err.setText('<font color="red">ERROR : </font>you have to fill from -> to Graduation Year for filtering ')
            self.set_dt()


# todo############################################### Rapport 3 #########################################################
class R3(QtWidgets.QWidget):
    def __init__(self, df=None, ar=False):
        super().__init__()
        uic.loadUi(os.path.join(os.path.dirname(__file__), 'ui/r3.ui'), self)
        self.df = df
        self.ar = ar
        # self.table.itemSelectionChanged.connect(self.table_select_event)
        self.table.clear()

        self.table.resizeColumnsToContents()
        self.from_txt.setValidator(QIntValidator())
        self.to_txt.setValidator(QIntValidator())
        for i in range(4):
            self.table.horizontalHeader().setSectionResizeMode(i, QHeaderView.Stretch)

        self.filter.clicked.connect(self.filtering)
        # self.filter.clicked.connect(self.filtering)
        self.produce.clicked.connect(self.export_to_exel)
        if self.ar:
            self.frame.setLayoutDirection(QtCore.Qt.RightToLeft)
            self.label_2.setText('من')
            self.label_3.setText('إلى')
            self.from_txt.setPlaceholderText('سنة')
            self.to_txt.setPlaceholderText('سنة')
            self.filter.setText('فلتر')
            self.produce.setText('إستخراج')
            self.label.setText("تصنيف حسب")
            self.comboBox.clear()
            self.comboBox.addItems(["الكليات", "التخصصات", "الأقسام"])

        self.df = self.df[self.df['Date of birth'] != '//']
        self.df['b_year'] = self.df['Date of birth'].str.split('/', expand=True)[0]
        self.df.loc[(self.df['b_year'] == '') | (self.df['b_year'] == 'nan') | (self.df['b_year'] == 'NaN'), 'b_year'] = \
        self.df['Graduation Year']
        self.df.loc[(self.df['Graduation Year'] == '') | (self.df['Graduation Year'] == 'nan') | (
                    self.df['Graduation Year'] == 'NaN'), 'Graduation Year'] = self.df['b_year']
        self.df['b_year'] = pd.to_numeric(self.df['b_year'], errors='coerce')
        self.df['age'] = self.df['Graduation Year'] - self.df['b_year']
        self.df['age'] = self.df['age'].fillna(0)
        self.df = self.df[self.df['age'] > 0]
        # self.df.loc[(self.df['age'] == 0.0) | (self.df['age'] == 'nan') | (self.df['age'] == 'NaN'), 'age'] = 20
        self.set_dt()

    def export_to_exel(self):
        if self.table.rowCount():
            data = []
            for r in range(self.table.rowCount()):
                rr = []
                for c in range(self.table.columnCount()):
                    rr.append(self.table.item(r, c).text() if self.table.item(r, c).text() else 0)
                data.append(rr)

            import arabic_reshaper
            from bidi.algorithm import get_display  # python-bidi

            self.plotdata = None

            self.plotdata = pd.DataFrame({
                get_display(arabic_reshaper.reshape("غلى الأقل")) if self.ar else "min": [float(i[1]) for i in data],
                get_display(arabic_reshaper.reshape("المتوسط")) if self.ar else "mean": [float(i[2]) for i in data],
                get_display(arabic_reshaper.reshape("على الأكثر")) if self.ar else "max": [float(i[3]) for i in data]},
                index=[get_display(arabic_reshaper.reshape(i[0])) for i in data]
            )

            self.plotdata.plot(kind="bar", figsize=(12, 7))
            plt.title(get_display(arabic_reshaper.reshape(self.title.text())) if self.ar else self.title.text())
            # plt.xlabel("Family Member")
            # plt.ylabel("Pies Consumed")
            plt.xticks(rotation=45, horizontalalignment='right')

            grapg_path = os.path.join(DESKTOP, 'graph.png')
            plt.savefig(grapg_path)

            # plt.show()
            document = Document()

            document.add_paragraph(self.title.text(),
                style='Intense Quote')

            document.add_picture(grapg_path, width=Inches(7))
            table = document.add_table(rows=1, cols=4)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = " "
            hdr_cells[1].text = "غلى الأقل" if self.ar else "min"
            hdr_cells[2].text = "المتوسط" if self.ar else "mean"
            hdr_cells[3].text = "على الأكثر" if self.ar else "max"

            for i in data:
                row_cells = table.add_row().cells
                for c in range(4):
                    row_cells[c].text = str(i[c])

            if self.table.rowCount():
                self.filename = QtWidgets.QFileDialog.getSaveFileName(caption='إستخراج' if self.ar else 'Export',
                                                                      filter="Word (*.doc *.docx)", directory=DESKTOP)[
                    0]
                if not QFileInfo(self.filename).suffix():
                    self.filename += '.docx'

            if self.filename:
                document.save(self.filename)
            print('done')
            os.remove(grapg_path)

    def set_dt(self, from_=None, to=None):
        G_years = list([i for i in self.df['Graduation Year']])
        if from_ is None or from_ < min(G_years) or from_ > max(G_years):
            from_ = min(G_years)
        if to is None or to > max(G_years) or to < min(G_years):
            to = max(G_years)
        self.headers = []
        self.rows = []
        self.dff = self.df[(self.df['Graduation Year'] >= from_) & (self.df['Graduation Year'] <= to)]

        if self.comboBox.currentIndex() == 0:
            l = list([i for i in self.dff['age']])
            ll = pd.Series(data=l)
            ll.dropna(inplace=True)

            self.rows = [
                ['الكلية' if self.ar else 'College', round(min(ll), 1), round(sum(ll) / len(ll), 1), round(max(ll), 1)]]
            self.headers = ['الكليات' if self.ar else 'Colleges', 'على الأقل' if self.ar else 'Min',
                            'المتوسط' if self.ar else 'Mean', 'على الأكثر' if self.ar else 'Max']

        else:
            key = self.comboBox.currentText()
            if self.ar:
                if self.comboBox.currentIndex() == 1:
                    key = 'Major'
                elif self.comboBox.currentIndex() == 2:
                    key = 'Department'
            elements = [i for i in set(list(self.dff[key]))]
            self.data = self.dff.groupby(key)
            self.headers = [self.comboBox.currentText(), 'على الأقل' if self.ar else 'Min',
                            'المتوسط' if self.ar else 'Mean', 'على الأكثر' if self.ar else 'Max']
            for elm in elements:
                l = [i for i in self.data.get_group(elm)['age']]
                ll = pd.Series(data=l)
                ll.dropna(inplace=True)

                self.rows.append([elm, round(min(ll), 1), round(sum(ll) / len(ll), 1), round(max(ll), 1)])
        msg = 'تصنيف أعمار الطلبة' if self.ar else f'min/mean/max of Age Grouping by {self.comboBox.currentText()}s from Graduation Year "{from_}" to "{to}"'
        self.title.setText(msg)
        self.table.setHorizontalHeaderLabels(self.headers)
        self.table.setColumnCount(len(self.headers))
        [self.table.removeRow(0) for _ in range(self.table.rowCount())]
        for r_n, r_d in enumerate(self.rows):
            self.table.insertRow(r_n)
            for c_n, d in enumerate(r_d):
                self.table.setItem(r_n, c_n, QtWidgets.QTableWidgetItem(str(d)))

        set0 = QtChart.QBarSet('على الأقل' if self.ar else 'Min')
        set1 = QtChart.QBarSet('المتوسط' if self.ar else 'Mean')
        set2 = QtChart.QBarSet('على الأكثر' if self.ar else 'Max')

        set0.append([i[1] for i in self.rows])
        set1.append([i[2] for i in self.rows])
        set2.append([i[3] for i in self.rows])

        series = QtChart.QBarSeries()
        series.append(set0)
        series.append(set1)
        series.append(set2)

        chart = QtChart.QChart()
        chart.addSeries(series)
        chart.setTitle(' ')
        chart.setAnimationOptions(QtChart.QChart.SeriesAnimations)

        axisX = QtChart.QBarCategoryAxis()
        # axisX.append(str(i) for i in range(1, len(data) + 1))
        axisX.append(str(i[0]) for i in self.rows)
        axisX.setLabelsAngle(90)
        axisX.setTitleText(f'{self.comboBox.currentText()}')
        font = QtGui.QFont()
        font.setPixelSize(5)
        axisX.tickFont = font
        all_v = []
        for i in self.rows:
            for x in i[1:]:
                all_v.append(x)
        axisY = QtChart.QValueAxis()
        axisY.setRange(0, max(all_v) if all_v else 0)
        axisY.setTitleText('سنة التخرج' if self.ar else "Graduation Year")

        chart.addAxis(axisX, Qt.AlignBottom)
        chart.addAxis(axisY, Qt.AlignLeft)

        chart.legend().setVisible(True)
        chart.legend().setAlignment(Qt.AlignBottom)
        chartView = QtChart.QChartView(chart)

        for i in reversed(range(self.graph_layout.count())):
            self.graph_layout.itemAt(i).widget().setParent(None)

        self.graph_layout.addWidget(chartView)

    def filtering(self):
        if self.from_txt.text() and self.to_txt.text():
            self.set_dt(int(self.from_txt.text()), int(self.to_txt.text()))
        else:
            # self.err.setText('<font color="red">ERROR : </font>you have to fill from -> to Graduation Year for filtering ')
            self.set_dt()


# todo############################################### Rapport 4 #########################################################
class R4(QtWidgets.QWidget):
    def __init__(self, df=None, ar=False):
        super().__init__()
        uic.loadUi(os.path.join(os.path.dirname(__file__), 'ui/r4.ui'), self)
        self.df = df
        self.ar = ar
        # self.table.itemSelectionChanged.connect(self.table_select_event)
        self.table.clear()

        self.table.resizeColumnsToContents()
        self.from_txt.setValidator(QIntValidator())
        self.to_txt.setValidator(QIntValidator())
        for i in range(4):
            self.table.horizontalHeader().setSectionResizeMode(i, QHeaderView.Stretch)

        if self.ar:
            self.frame.setLayoutDirection(QtCore.Qt.RightToLeft)
            self.label_2.setText('من')
            self.label_3.setText('إلى')
            self.from_txt.setPlaceholderText('سنة')
            self.to_txt.setPlaceholderText('سنة')
            self.filter.setText('فلتر')
            self.produce.setText('إستخراج')
            self.label.setText("تصنيف حسب")
            self.comboBox.clear()
            self.comboBox.addItems(["الكليات", "التخصصات", "الأقسام"])

        self.filter.clicked.connect(self.filtering)
        # self.filter.clicked.connect(self.filtering)
        self.produce.clicked.connect(self.export_to_exel)

        self.df['GPA'] = self.df['GPA'].fillna(0)
        self.df = self.df[self.df['GPA'] > 0]
        # self.df.loc[(self.df['age'] == 0.0) | (self.df['age'] == 'nan') | (self.df['age'] == 'NaN'), 'age'] = 20
        self.set_dt()

    def export_to_exel(self):
        if self.table.rowCount():

            data = []
            for r in range(self.table.rowCount()):
                rr = []
                for c in range(self.table.columnCount()):
                    rr.append(self.table.item(r, c).text() if self.table.item(r, c).text() else 0)
                data.append(rr)

            import arabic_reshaper
            from bidi.algorithm import get_display  # python-bidi

            self.plotdata = None

            self.plotdata = pd.DataFrame({
                get_display(arabic_reshaper.reshape("على الأقل")) if self.ar else "min": [float(i[1]) for i in data],
                get_display(arabic_reshaper.reshape("المتوسط")) if self.ar else "mean" : [float(i[2]) for i in data],
                get_display(arabic_reshaper.reshape("على الأكثر")) if self.ar else "max" : [float(i[3]) for i in data]},
                index=[get_display(arabic_reshaper.reshape(i[0])) for i in data]
            )

            self.plotdata.plot(kind="bar", figsize=(12, 7))
            plt.title(get_display(arabic_reshaper.reshape(self.title.text())) if self.ar else self.title.text())
            # plt.xlabel("Family Member")
            # plt.ylabel("Pies Consumed")
            plt.xticks(rotation=45, horizontalalignment='right')

            grapg_path = os.path.join(DESKTOP, 'graph.png')
            plt.savefig(grapg_path)

            # plt.show()
            document = Document()

            document.add_paragraph(self.title.text(),
                style='Intense Quote')

            document.add_picture(grapg_path, width=Inches(7))
            table = document.add_table(rows=1, cols=4)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = " - "
            hdr_cells[1].text = "على الأقل" if self.ar else "min"
            hdr_cells[2].text = "المتوسط" if self.ar else "mean"
            hdr_cells[3].text = "على الأكثر" if self.ar else  "max"

            for i in data:
                row_cells = table.add_row().cells
                for c in range(4):
                    row_cells[c].text = str(i[c])

            if self.table.rowCount():
                self.filename = QtWidgets.QFileDialog.getSaveFileName(caption='إستخراج' if self.ar else 'Export',
                                                                      filter="Word (*.doc *.docx)", directory=DESKTOP)[
                    0]
                if not QFileInfo(self.filename).suffix():
                    self.filename += '.docx'

            if self.filename:
                document.save(self.filename)
            print('done')
            os.remove(grapg_path)

    def set_dt(self, from_=None, to=None):
        G_years = list([i for i in self.df['Graduation Year']])
        if from_ is None or from_ < min(G_years) or from_ > max(G_years):
            from_ = min(G_years)
        if to is None or to > max(G_years) or to < min(G_years):
            to = max(G_years)
        self.headers = []
        self.rows = []
        self.dff = self.df[(self.df['Graduation Year'] >= from_) & (self.df['Graduation Year'] <= to)]

        if self.comboBox.currentIndex() == 0:
            l = list([i for i in self.dff['GPA']])
            ll = pd.Series(data=l)
            ll.dropna(inplace=True)

            self.rows = [
                ['الكلية' if self.ar else 'College', round(min(ll), 1), round(sum(ll) / len(ll), 1), round(max(ll), 1)]]
            self.headers = ['الكليات' if self.ar else 'Colleges', 'على الأقل' if self.ar else 'Min',
                            'المتوسط' if self.ar else 'Mean', 'على الأكثر' if self.ar else 'Max']

        else:
            key = self.comboBox.currentText()
            if self.ar:
                if self.comboBox.currentIndex() == 1:
                    key = 'Major'
                elif self.comboBox.currentIndex() == 2:
                    key = 'Department'
            elements = [i for i in set(list(self.dff[key]))]
            self.data = self.dff.groupby(key)
            self.headers = [self.comboBox.currentText(), 'على الأقل' if self.ar else 'Min',
                            'المتوسط' if self.ar else 'Mean', 'على الأكثر' if self.ar else 'Max']
            for elm in elements:
                l = [i for i in self.data.get_group(elm)['GPA']]
                ll = pd.Series(data=l)
                ll.dropna(inplace=True)
                self.rows.append([elm, round(min(ll), 1), round(sum(ll) / len(ll), 1), round(max(ll), 1)])
        ttl = "تصنيف الحد الأدنى / المتوسط ​​/ الحد الأقصى لمتوسط ​​الدرجة حسب سنوات التخرج" if self.ar else f'min/mean/max of GPA Grouping by {self.comboBox.currentText()}s from Graduation Year "{from_}" to "{to}"'
        self.title.setText(ttl)
        self.table.setHorizontalHeaderLabels(self.headers)
        self.table.setColumnCount(len(self.headers))
        [self.table.removeRow(0) for _ in range(self.table.rowCount())]
        for r_n, r_d in enumerate(self.rows):
            self.table.insertRow(r_n)
            for c_n, d in enumerate(r_d):
                self.table.setItem(r_n, c_n, QtWidgets.QTableWidgetItem(str(d)))

        set0 = QtChart.QBarSet('على الأقل' if self.ar else 'Min')
        set1 = QtChart.QBarSet('المتوسط' if self.ar else 'Mean')
        set2 = QtChart.QBarSet('على الأكثر' if self.ar else 'Max')

        set0.append([i[1] for i in self.rows])
        set1.append([i[2] for i in self.rows])
        set2.append([i[3] for i in self.rows])

        series = QtChart.QBarSeries()
        series.append(set0)
        series.append(set1)
        series.append(set2)

        chart = QtChart.QChart()
        chart.addSeries(series)
        chart.setTitle(' ')
        chart.setAnimationOptions(QtChart.QChart.SeriesAnimations)

        axisX = QtChart.QBarCategoryAxis()
        # axisX.append(str(i) for i in range(1, len(data) + 1))
        axisX.append(str(i[0]) for i in self.rows)
        axisX.setLabelsAngle(90)
        axisX.setTitleText(f'{self.comboBox.currentText()}')
        font = QtGui.QFont()
        font.setPixelSize(5)
        axisX.tickFont = font
        all_v = []
        for i in self.rows:
            for x in i[1:]:
                all_v.append(x)
        axisY = QtChart.QValueAxis()
        axisY.setRange(0, max(all_v) if all_v else 0)
        axisY.setTitleText("سنوات التخرج" if self.ar else "Graduation Year")

        chart.addAxis(axisX, Qt.AlignBottom)
        chart.addAxis(axisY, Qt.AlignLeft)

        chart.legend().setVisible(True)
        chart.legend().setAlignment(Qt.AlignBottom)
        chartView = QtChart.QChartView(chart)

        for i in reversed(range(self.graph_layout.count())):
            self.graph_layout.itemAt(i).widget().setParent(None)

        self.graph_layout.addWidget(chartView)

    def filtering(self):
        if self.from_txt.text() and self.to_txt.text():
            self.set_dt(int(self.from_txt.text()), int(self.to_txt.text()))
        else:
            # self.err.setText('<font color="red">ERROR : </font>you have to fill from -> to Graduation Year for filtering ')
            self.set_dt()


# todo############################################### Rapport 5 #########################################################
class R5(QtWidgets.QWidget):
    def __init__(self, df=None, ar=False):
        super().__init__()
        uic.loadUi(os.path.join(os.path.dirname(__file__), 'ui/r5.ui'), self)
        self.df = df
        self.ar = ar
        # self.table.itemSelectionChanged.connect(self.table_select_event)
        self.table.clear()
        self.table_header = ['المناطق' if self.ar else "Regions", 'عدد الطلاب' if self.ar else 'Number of students']
        self.table.setColumnCount(len(self.table_header))
        self.table.setHorizontalHeaderLabels(self.table_header)
        self.table.resizeColumnsToContents()
        self.from_txt.setValidator(QIntValidator())
        self.to_txt.setValidator(QIntValidator())
        for i in range(len(self.table_header)):
            self.table.horizontalHeader().setSectionResizeMode(i, QHeaderView.Stretch)

        self.df = self.df.drop('Major', 1)
        if self.ar:
            self.frame.setLayoutDirection(QtCore.Qt.RightToLeft)
            self.label_2.setText('من')
            self.label_3.setText('إلى')
            self.from_txt.setPlaceholderText('سنة')
            self.to_txt.setPlaceholderText('سنة')
            self.filter.setText('فلتر')
            self.produce.setText('إستخراج')
        self.set_dt()
        self.filter.clicked.connect(self.filtering)
        self.produce.clicked.connect(self.export_to_exel)

    def export_to_exel(self):
        if self.table.rowCount():

            data = []
            for r in range(self.table.rowCount()):
                rr = []
                for c in range(self.table.columnCount()):
                    rr.append(self.table.item(r, c).text() if self.table.item(r, c).text() else 0)
                data.append(rr)

            import arabic_reshaper
            from bidi.algorithm import get_display  # python-bidi

            self.plotdata = None

            self.plotdata = pd.DataFrame({get_display(arabic_reshaper.reshape('عدد الطلاب')) if self.ar else 'Number of students': [float(i[1]) for i in data]},
                                         index=[get_display(arabic_reshaper.reshape(i[0])) for i in data]
                                         )

            self.plotdata.plot(kind="bar", figsize=(12, 7))
            plt.title(get_display(arabic_reshaper.reshape(self.title_.text())))
            # plt.xlabel("Family Member")
            # plt.ylabel("Pies Consumed")
            plt.xticks(rotation=45, horizontalalignment='right')

            grapg_path = os.path.join(DESKTOP, 'graph.png')
            plt.savefig(grapg_path)

            # plt.show()
            document = Document()

            document.add_paragraph(self.title_.text(), style='Intense Quote')

            document.add_picture(grapg_path, width=Inches(7))
            table = document.add_table(rows=1, cols=2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'المناطق' if self.ar else "Regions"
            hdr_cells[1].text = 'عدد الطلاب' if self.ar else 'Number of students'

            for i in data:
                row_cells = table.add_row().cells
                for c in range(2):
                    row_cells[c].text = str(i[c])

            if self.table.rowCount():
                self.filename = QtWidgets.QFileDialog.getSaveFileName(caption='إستخراج' if self.ar else 'Export',
                                                                      filter="Word (*.doc *.docx)", directory=DESKTOP)[
                    0]
                if not QFileInfo(self.filename).suffix():
                    self.filename += '.docx'

            if self.filename:
                document.save(self.filename)
            print('done')
            os.remove(grapg_path)

    def filtering(self):
        if self.from_txt.text() and self.to_txt.text():
            self.set_dt(int(self.from_txt.text()), int(self.to_txt.text()))
        else:
            # self.err.setText('<font color="red">ERROR : </font>you have to fill from -> to Graduation Year for filtering ')
            self.set_dt()

    def set_dt(self, from_=None, to=None):
        G_years = list([i for i in self.df['Graduation Year']])
        if from_ is None or from_ < min(G_years) or from_ > max(G_years):
            from_ = min(G_years)
        if to is None or to > max(G_years) or to < min(G_years):
            to = max(G_years)
        msg = 'تصنيف عدد الطلبة حسب الجهات' if self.ar else f'Number Students Grouping by Regions from Graduation Year "{from_}" to "{to}"'
        self.title_.setText(msg)
        self.new_df = self.df[(self.df['Graduation Year'] >= from_) & (self.df['Graduation Year'] <= to)]
        rows = len(self.new_df)
        if rows > 0:
            p1 = ['Riyadh', 'Thadiq', 'Shaqra', 'Al-Kharj', 'Huraymila', 'Az Zulfi', 'Al Duwadimi', 'Afif', 'Aflaj',
                  'As Sulayyil', 'Al Majmah', 'Howtat Bani Tamim', 'Al Quwaiiyah', 'Wadi Al-Dawasir', 'Rumah',
                  'Al Ghat',
                  'Hautat Sudair', 'Al Uyaynah', 'Al Artawiyah', 'Ar Rayn', 'Al Hariq', 'Dhurma', 'Al-Muzahmiya',
                  'Sajir',
                  'Ad Diriyah', 'الرياض', 'ثادق', 'شقراء', 'الخرج', 'حريملاء', 'الزلفي', 'الدوادمي', 'عفيف', 'الافلاج',
                  'السليل',
                  'المجمعة', 'حوطة بني تميم', 'القويعية', 'وادي الدواسر', 'رماح', 'الغاط', 'حوطه سدير', 'العيينة',
                  'الأرطاوية', 'الرين',
                  'الحريق', 'ضرماء', 'المزاحمية', 'ساجر', 'الدرعيه']

            p2 = ['Taif', 'Ranyah', 'Makkah', 'Jeddah', 'Turbah', 'Al Qunfudhah', 'Al Khurma', 'Ardiya Al Janubia',
                  'الطائف',
                  'رنيه',
                  'مكة المكرمة',
                  'جدة',
                  'تربة',
                  'القنفذة',
                  'الخرمة',
                  'العرضية الجنوبية',
                  ]

            p3 = [
                'Medinah',
                'Al Mahd',
                'Yanbu',
                'Al Henakiyah',
                'AlUla',
                'Khaybar',
                'Badr',
                'المدينة المنورة',
                'المهد',
                'ينبع',
                'الحناكية',
                'العلا',
                'خيبر',
                'بدر',

            ]

            p4 = [
                'Al Bukayriyah',
                'Buraydah',
                'Unayzah',
                'Ar Rass',
                'Al Mithnab',
                'Al Badayea',
                'البكيرية',
                'بريدة',
                'عنيزة',
                'الرس',
                'المذنب',
                'البدائع'

            ]

            p5 = [
                'Al Jubail',
                'Al Khobar',
                'Khafji',
                'Hafr Al - Batin',
                'Al Ahsa',
                'Dammam',
                'Nairyah',
                'Al Hofuf',
                'Qatif',
                'Dhahran',
                'Qaryat Al Ulya',
                'الجبيل',
                'الخبر',
                'الخفجي',
                'حفر الباطن',
                'الاحساء',
                'الدمام',
                'النعيريه', 'الهفوف', 'القطيف', 'قرية',

            ]

            p6 = [
                'Al Namas',
                'Bisha',
                'Sarat Abidah',
                'Dhahran Al Janub',
                'Khamis Mushait',
                'Abha',
                'Mahayel Aseer',
                'Khamis Mushait',
                'Aseer',
                'Ahad Rafidah',
                'Balqarn',
                'Ragal Almaa',
                'Tathleeth',
                'النماص',
                'بيشه',
                'سراة عبيدة',
                'ظهران الجنوب',
                'خميس مشيط',
                'أبها',
                'محايل', 'عسير', 'خميس', 'مشيط', 'عسير', 'أحدرفيدة', 'بلقرن', 'رجال', 'المع', 'تثليث'

            ]

            p7 = [
                'Haql',
                'Umluj',
                'Tabuk',
                'Duba',
                'Tayma',
                'Al Wajh',
                'حقل',
                'املج',
                'تبوك',
                'ضباء',
                'تيماء',
                'الوجه'

            ]

            p8 = [
                'Hail',
                'Baqaa',
                'حائل',
                'بقعاء'

            ]

            p9 = [
                'Rafha',
                'Arar',
                'Al Qurayyat',
                'Turaif',
                'رفحاء',
                'عرعر',
                'القريات',
                'طريف'

            ]

            p10 = [
                'Sabya',
                'Fayfa',
                'Addayer',
                'Abu Arish',
                'Samtah',
                'Farasan',
                'Jazan',
                'صبيا',
                'فيفا',
                'الداير', 'أبو عريش',
                'صامطه',
                'فرسان',
                'جيزان'

            ]

            p11 = [
                'Najran',
                'Sharorah',
                'نجران',
                'شروره'

            ]

            p12 = [
                'Almandaq',
                'Al Bahah',
                'Baljurashi',
                'المندق',
                'الباحة',
                'بلجرشي'

            ]

            p13 = [
                'Dumah Al Jandal',
                'Tubarjal',
                'Al Jouf',
                'دومة الجندل',
                'طبرجل',
                'الجوف'

            ]

            p14 = [
                'Sana',
                'Kuwait',
                'Muscat',
                'صنعاء',
                'الكويت',
                'مسقط'

            ]

            regions_ = [['منطقة الرياض', 'Riyadh'], ['منطقة مكة المكرمة', 'Makkah'],[ 'المدينة المنورة', 'Madinah'], ['منطقة القصيم', 'Qassim'], ['الشرقية', 'Eastern Province'], ['منطقة عسير', 'Asir'],
                        ['منطقة تبوك', 'Tabuk'], ['حائل', 'Hail'], ['منطقة الحدود الشمالية', 'Northern Borders'], ['جازان', 'Jazan'], ['منطقة نجران', 'Najran'],
                        ['منطقة الباحة', 'Al Bahah'], ['منطقة الجوف', 'Al Jouf'], ['طلاب دوليين', 'international']
                        ]
            rr = []
            for row in self.new_df['Place of issue']:
                if row in p1:
                    rr.append(regions_[0][0] if self.ar else regions_[0][1])
                elif row in p2:
                    rr.append(regions_[1][0] if self.ar else regions_[1][1])
                elif row in p3:
                    rr.append(regions_[2][0] if self.ar else regions_[2][1])
                elif row in p4:
                    rr.append(regions_[3][0] if self.ar else regions_[3][1])
                elif row in p5:
                    rr.append(regions_[4][0] if self.ar else regions_[4][1])
                elif row in p6:
                    rr.append(regions_[5][0] if self.ar else regions_[5][1])
                elif row in p7:
                    rr.append(regions_[6][0] if self.ar else regions_[6][1])
                elif row in p8:
                    rr.append(regions_[7][0] if self.ar else regions_[7][1])
                elif row in p9:
                    rr.append(regions_[8][0] if self.ar else regions_[8][1])
                elif row in p10:
                    rr.append(regions_[9][0] if self.ar else regions_[9][1])
                elif row in p11:
                    rr.append(regions_[10][0] if self.ar else regions_[10][1])
                elif row in p12:
                    rr.append(regions_[11][0] if self.ar else regions_[11][1])
                elif row in p13:
                    rr.append(regions_[12][0] if self.ar else regions_[12][1])
                elif row in p14:
                    rr.append(regions_[13][0] if self.ar else regions_[13][1])
                else:
                    rr.append(row)

            self.new_df['region'] = rr

            places = pd.Series(data=[i for i in set(list(self.new_df['region']))])
            places.dropna(inplace=True)
            gk = self.new_df.groupby('region')
            data = []
            for re in places:
                data.append([re, len(gk.get_group(re)['Student ID'])])
            # [print(i) for i in data]
            [self.table.removeRow(0) for _ in range(self.table.rowCount())]
            for r_n, r_d in enumerate(data):
                self.table.insertRow(r_n)
                for c_n, d in enumerate(r_d):
                    self.table.setItem(r_n, c_n, QtWidgets.QTableWidgetItem(str(d)))

            set0 = QtChart.QBarSet('عدد الطلاب' if self.ar else 'Number of students')

            set0.append([i[1] for i in data])

            series = QtChart.QBarSeries()
            series.append(set0)

            chart = QtChart.QChart()
            chart.addSeries(series)
            chart.setTitle(' ')
            chart.setAnimationOptions(QtChart.QChart.SeriesAnimations)

            axisX = QtChart.QBarCategoryAxis()
            # axisX.append(str(i) for i in range(1, len(data) + 1))
            axisX.append(str(i[0]) for i in data)
            axisX.setLabelsAngle(90)
            axisX.setTitleText('المناطق' if self.ar else "Regions")
            font = QtGui.QFont()
            font.setPixelSize(5)
            axisX.tickFont = font

            all_v = []
            for i in data:
                for x in i[1:]:
                    all_v.append(x)
            axisY = QtChart.QValueAxis()
            axisY.setRange(0, max(all_v) if all_v else 0)
            # axisY.setTitleText("Years in College")

            chart.addAxis(axisX, Qt.AlignBottom)
            chart.addAxis(axisY, Qt.AlignLeft)

            chart.legend().setVisible(True)
            chart.legend().setAlignment(Qt.AlignBottom)
            chartView = QtChart.QChartView(chart)

            for i in reversed(range(self.verticalLayout_3.count())):
                self.verticalLayout_3.itemAt(i).widget().setParent(None)

            self.verticalLayout_3.addWidget(self.frame_2)
            for i in reversed(range(self.graph_layout.count())):
                self.graph_layout.itemAt(i).widget().setParent(None)

            self.graph_layout.addWidget(chartView)
        else:
            lbl = Err()
            for i in reversed(range(self.verticalLayout_3.count())):
                self.verticalLayout_3.itemAt(i).widget().setParent(None)
            self.verticalLayout_3.addWidget(lbl)


# todo############################################### Rapport 6 #########################################################
class R6(QtWidgets.QWidget):
    def __init__(self, df=None, ar=False):
        super().__init__()
        uic.loadUi(os.path.join(os.path.dirname(__file__), 'ui/r6.ui'), self)
        self.df = df
        self.ar = ar
        # self.table.itemSelectionChanged.connect(self.table_select_event)

        self.table_header = []
        self.from_txt.setValidator(QIntValidator())
        self.to_txt.setValidator(QIntValidator())

        self.df = self.df.drop('Major', 1)
        if self.ar:
            self.frame.setLayoutDirection(QtCore.Qt.RightToLeft)
            self.label_2.setText('من')
            self.label_3.setText('إلى')
            self.from_txt.setPlaceholderText('سنة')
            self.to_txt.setPlaceholderText('سنة')
            self.filter.setText('فلتر')
            self.produce.setText('إستخراج')
            self.checkBox.setText('متوسط الدرجة')
        self.set_dt()
        self.filter.clicked.connect(self.filtering)
        self.produce.clicked.connect(self.export_to_exel)

    def export_to_exel(self):
        if self.table.rowCount():

            data = []

            for r in range(self.table.rowCount()):
                rr = []
                for c in range(self.table.columnCount()):
                    rr.append(self.table.item(r, c).text() if self.table.item(r, c).text() else 0)
                data.append(rr)

            import arabic_reshaper
            from bidi.algorithm import get_display  # python-bidi

            self.plotdata = None
            if self.checkBox.isChecked():
                self.plotdata = pd.DataFrame({
                    get_display(arabic_reshaper.reshape('غلى الأقل')) if self.ar else "min": [float(i[1]) for i in data],
                    get_display(arabic_reshaper.reshape('المتوسط')) if self.ar else "mean": [float(i[2]) for i in data],
                    get_display(arabic_reshaper.reshape('غلى الأكثر')) if self.ar else "max": [float(i[3]) for i in data]},
                    index=[get_display(arabic_reshaper.reshape(i[0])) for i in data]
                )
            else:
                self.plotdata = pd.DataFrame({
                    get_display(arabic_reshaper.reshape('عدد الطلاب')) if self.ar else 'Number of students': [float(i[1]) for i in data]},
                    index=[get_display(arabic_reshaper.reshape(i[0])) for i in data]
                )

            self.plotdata.plot(kind="bar", figsize=(12, 7))
            plt.title(get_display(arabic_reshaper.reshape(self.title_.text())) if self.ar else self.title_.text())
            # plt.xlabel("Family Member")
            # plt.ylabel("Pies Consumed")
            plt.xticks(rotation=45, horizontalalignment='right')

            grapg_path = os.path.join(DESKTOP, 'graph.png')
            plt.savefig(grapg_path)

            # plt.show()
            document = Document()

            document.add_paragraph(self.title_.text(), style='Intense Quote')

            document.add_picture(grapg_path, width=Inches(7))
            table = document.add_table(rows=1, cols= 4 if self.checkBox.isChecked() else 2)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'أصلي / غير أصلي' if self.ar else  "Native/non-Native"
            hdr_cells[1].text = ('على الأقل' if self.ar else "min" )if self.checkBox.isChecked() else ('عدد الطلاب' if self.ar else 'Number of students')
            if self.checkBox.isChecked():
                hdr_cells[2].text = "المتوسط" if self.ar else "mean"
                hdr_cells[3].text = "على الأ كثر" if self.ar else "max"

            for i in data:
                row_cells = table.add_row().cells
                for c in range(4 if self.checkBox.isChecked() else 2):
                    row_cells[c].text = str(i[c])

            if self.table.rowCount():
                self.filename = QtWidgets.QFileDialog.getSaveFileName(caption='إستخراج' if self.ar else 'Export',
                                                                      filter="Word (*.doc *.docx)", directory=DESKTOP)[
                    0]
                if not QFileInfo(self.filename).suffix():
                    self.filename += '.docx'

            if self.filename:
                document.save(self.filename)
            print('done')
            os.remove(grapg_path)

    def filtering(self):

        if self.from_txt.text() and self.to_txt.text():
            self.set_dt(int(self.from_txt.text()), int(self.to_txt.text()))

        else:
            # self.err.setText('<font color="red">ERROR : </font>you have to fill from -> to Graduation Year for filtering ')
            self.set_dt()

    def set_dt(self, from_=None, to=None):
        gpa = self.checkBox.isChecked()
        G_years = list([i for i in self.df['Graduation Year']])
        if from_ is None or from_ < min(G_years) or from_ > max(G_years):
            from_ = min(G_years)
        if to is None or to > max(G_years) or to < min(G_years):
            to = max(G_years)
        ttl = ''
        if gpa:
            ttl = "تصنيف الحد الأدنى / المتوسط / الحد الأقصى لمتوسط الدرجة" if self.ar else f'min/mean/max of GPA Grouping by Native/non-Native from Graduation Year "{from_}" to "{to}"'
        else:
            arr = "عدد تجمع الطلاب من السكان الأصليين / غير الأصليين حسب المنطقة"
            ttl = arr if self.ar else f'Number of Native/non-Native Students Grouping by Origin from Graduation Year "{from_}" to "{to}"'
        self.title_.setText(ttl)
        self.df = self.df[self.df['Type of ID'] != '']
        self.new_df = self.df[(self.df['Graduation Year'] >= from_) & (self.df['Graduation Year'] <= to)]
        self.new_df.loc[self.new_df['Type of ID'] == 'بطاقة أحوال', 'origin'] = 'native'
        self.new_df.loc[self.new_df['Type of ID'] != 'بطاقة أحوال', 'origin'] = 'non-native'

        rows = len(self.new_df)
        if rows > 0:

            gk = self.new_df.groupby('origin')
            self.data = []
            if gpa:
                for re in ['native', 'non-native']:
                    ll = [i for i in gk.get_group(re)['GPA']]
                    ll = pd.Series(ll)
                    ll.dropna(inplace=True)
                    u = re
                    if self.ar:
                        if re == 'native' : u = "أصلي"
                        else : u = "أجنبي"
                    self.data.append([u, min(ll), round(sum(ll) / len(ll), 1), max(ll)])
                self.table_header = ["أصلي/أجنبي" if self.ar else 'Native/non-native', "على الأقل" if self.ar else 'min', "المتوسط" if self.ar else 'mean', "على الأكثر" if self.ar else 'max']
            else:
                for re in ['native', 'non-native']:
                    u = re
                    if self.ar:
                        if re == 'native':
                            u = "أصلي"
                        else:
                            u = "أجنبي"
                    self.data.append([u, len([i for i in gk.get_group(re)['Student ID']])])

                self.table_header = ["أصلي/أجنبي" if self.ar else 'Native/non-native', "عدد الطلاب" if self.ar else 'Number of Students']

            self.table.clear()
            # [print(i) for i in data]

            self.table.setColumnCount(len(self.table_header))
            for i in range(len(self.table_header)):
                self.table.horizontalHeader().setSectionResizeMode(i, QHeaderView.Stretch)
            #
            self.table.setHorizontalHeaderLabels(self.table_header)
            self.table.resizeColumnsToContents()

            if self.table.rowCount():
                [self.table.removeRow(0) for _ in range(self.table.rowCount())]

            # [print(i) for i in data]
            for r_n, r_d in enumerate(self.data):
                self.table.insertRow(r_n)
                for c_n, d in enumerate(r_d):
                    self.table.setItem(r_n, c_n, QtWidgets.QTableWidgetItem(str(d)))

            if gpa:
                set0 = QtChart.QBarSet( "على الأقل" if self.ar else 'min')
                set1 = QtChart.QBarSet("المتوسط" if self.ar else 'mean')
                set2 = QtChart.QBarSet( "على الأكثر" if self.ar else 'max')

                set0.append([i[1] for i in self.data])
                set1.append([i[2] for i in self.data])
                set2.append([i[3] for i in self.data])

                series = QtChart.QBarSeries()
                series.append(set0)
                series.append(set1)
                series.append(set2)
            else:
                set0 = QtChart.QBarSet( "عدد الطلاب" if self.ar else 'Number of Students')
                set0.append([i[1] for i in self.data])
                series = QtChart.QBarSeries()
                series.append(set0)

            chart = QtChart.QChart()
            chart.addSeries(series)
            chart.setTitle(' ')
            chart.setAnimationOptions(QtChart.QChart.SeriesAnimations)

            axisX = QtChart.QBarCategoryAxis()
            # axisX.append(str(i) for i in range(1, len(data) + 1))
            axisX.append(str(i[0]) for i in self.data)

            # axisX.setLabelsAngle(90)
            # axisX.setTitleText("Regions")
            font = QtGui.QFont()
            font.setPixelSize(5)
            axisX.tickFont = font

            all_v = []
            for i in self.data:
                for x in i[1:]:
                    all_v.append(x)
            axisY = QtChart.QValueAxis()
            axisY.setRange(0, max(all_v) if all_v else 0)
            # axisY.setTitleText("Years in College")

            chart.addAxis(axisX, Qt.AlignBottom)
            chart.addAxis(axisY, Qt.AlignLeft)

            chart.legend().setVisible(True)
            chart.legend().setAlignment(Qt.AlignBottom)
            chartView = QtChart.QChartView(chart)

            for i in reversed(range(self.verticalLayout_3.count())):
                self.verticalLayout_3.itemAt(i).widget().setParent(None)

            self.verticalLayout_3.addWidget(self.frame_2)
            for i in reversed(range(self.graph_layout.count())):
                self.graph_layout.itemAt(i).widget().setParent(None)

            self.graph_layout.addWidget(chartView)
        else:
            lbl = Err()
            for i in reversed(range(self.verticalLayout_3.count())):
                self.verticalLayout_3.itemAt(i).widget().setParent(None)
            self.verticalLayout_3.addWidget(lbl)


class Err(QtWidgets.QFrame):
    def __init__(self):
        super().__init__()
        uic.loadUi(os.path.join(os.path.dirname(__file__), 'ui/err.ui'), self)

        self.label.setPixmap(QPixmap('src/nodt.png'))
        self.label.setScaledContents(True)


class Loading(QtWidgets.QWidget):
    def __init__(self, gif='loading.gif'):
        super().__init__()
        uic.loadUi(os.path.join(os.path.dirname(__file__), 'ui/loading.ui'), self)
        self.gif = QtGui.QMovie(f'src/{gif}')
        self.label.setMovie(self.gif)
        self.gif.start()


if __name__ == '__main__':
    app = QtWidgets.QApplication(sys.argv)
    splash = Splash()
    splash.show()
    sys.exit(app.exec_())
